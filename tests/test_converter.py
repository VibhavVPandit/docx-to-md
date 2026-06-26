"""Automated checks for the docx -> markdown converter.

Each test builds a small .docx in a temp directory (via the fixture builders),
converts it, and asserts the markdown contains the expected structure.
"""

import os
import sys

import pytest

_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)                       # tests/ -> generate_fixtures
sys.path.insert(0, os.path.dirname(_HERE))      # project root -> docx_to_md

from docx_to_md import docx_to_md  # noqa: E402
import generate_fixtures as gf      # noqa: E402


def convert(builder, tmp_path, **kwargs):
    """Build a doc with *builder*, save it, convert it, return (md, warnings)."""
    docx_path = os.path.join(str(tmp_path), "in.docx")
    out_path = os.path.join(str(tmp_path), "out.md")
    builder().save(docx_path)
    return docx_to_md(docx_path, out_path, **kwargs)


def test_inline_formatting(tmp_path):
    md, _ = convert(gf.build_inline_formatting, tmp_path)
    assert "**bold**" in md
    assert "*italic*" in md
    assert "***bold italic***" in md
    assert "~~struck out~~" in md
    # underlined text survives as plain text (markdown has no underline).
    assert "underlined" in md


def test_headings(tmp_path):
    md, _ = convert(gf.build_headings, tmp_path)
    assert "# Heading Level 1" in md
    assert "## Heading Level 2" in md
    assert "### Heading Level 3" in md
    assert "#### Heading Level 4" in md
    assert "Body text under heading 1." in md


def test_lists(tmp_path):
    md, _ = convert(gf.build_lists, tmp_path)
    assert "- First bullet" in md
    assert "  - Nested bullet" in md
    assert "- Second bullet" in md
    assert "1. Step one" in md
    assert "1. Step two" in md


def test_simple_table(tmp_path):
    md, _ = convert(gf.build_simple_table, tmp_path)
    assert "| Name | Role | City |" in md
    assert "| --- | --- | --- |" in md
    assert "| Ada | Engineer | London |" in md
    assert "<table>" not in md


def test_formatted_table(tmp_path):
    md, _ = convert(gf.build_formatted_table, tmp_path)
    assert "**Feature**" in md
    assert "**Status**" in md
    assert "*Export*" in md
    assert "**done**" in md


def test_merged_table_falls_back_to_html(tmp_path):
    md, warnings = convert(gf.build_merged_table, tmp_path)
    assert "<table>" in md
    assert 'colspan="2"' in md
    assert "Quarterly Results" in md
    assert "<!-- note:" in md
    assert any("HTML" in w for w in warnings)


def test_links_and_code(tmp_path):
    md, _ = convert(gf.build_links_and_code, tmp_path)
    assert "[Example](https://example.com)" in md
    assert "`pip install docx`" in md


def test_embedded_image(tmp_path):
    md, _ = convert(gf.build_with_image, tmp_path)
    assert "![" in md and "](" in md
    assert ".png)" in md
    images_dir = os.path.join(str(tmp_path), "images")
    assert os.path.isdir(images_dir)
    assert len(os.listdir(images_dir)) == 1


def test_kitchen_sink(tmp_path):
    md, _ = convert(gf.build_kitchen_sink, tmp_path)
    assert "# Project Report" in md
    assert "## Highlights" in md
    assert "- Shipped feature" in md
    assert "  - Sub-detail" in md
    assert "1. Next milestone" in md
    assert "[spec](https://example.com/spec)" in md
    assert "`config`" in md
    assert "**Important:**" in md or "**Important: **" in md
    assert "| **Metric** | **Value** |" in md  # header cells are bold
    assert "*99.9%*" in md


def test_force_html_tables(tmp_path):
    md, _ = convert(gf.build_simple_table, tmp_path, force_html_tables=True)
    assert "<table>" in md
    assert "| Name |" not in md


def test_pipe_escaped_in_cells(tmp_path):
    from docx import Document

    def builder():
        doc = Document()
        t = doc.add_table(rows=1, cols=1)
        t.rows[0].cells[0].text = "a | b"
        return doc

    md, _ = convert(builder, tmp_path)
    assert "a \\| b" in md


def test_special_chars_escaped(tmp_path):
    from docx import Document

    def builder():
        doc = Document()
        doc.add_paragraph("use *literal* and _under_ here")
        return doc

    md, _ = convert(builder, tmp_path)
    assert "\\*literal\\*" in md
    assert "\\_under\\_" in md


def _monospace_doc():
    """A document set entirely in a monospace font, with one bold run."""
    from docx import Document

    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("START CLERICAL ENC ")
    r1.font.name = "Consolas"
    r2 = p.add_run("bold word")
    r2.bold = True
    r2.font.name = "Consolas"
    r3 = p.add_run(" then more log text.")
    r3.font.name = "Consolas"
    return doc


def test_monospace_document_not_wrapped_as_code(tmp_path):
    # A whole-document monospace font must NOT turn every run into a code span;
    # bold should survive as **bold**, not get swallowed by backticks.
    md, warnings = convert(_monospace_doc, tmp_path)
    assert "`" not in md
    assert "**bold word**" in md
    assert any("code" in w for w in warnings)


def test_code_detection_can_be_forced_on(tmp_path):
    md, _ = convert(_monospace_doc, tmp_path, code_detection="on")
    assert "`" in md


def test_inline_code_still_detected_in_normal_doc(tmp_path):
    # A monospace run that is the exception (not the body font) stays code.
    md, _ = convert(gf.build_links_and_code, tmp_path)
    assert "`pip install docx`" in md


def test_output_ends_with_newline(tmp_path):
    md, _ = convert(gf.build_headings, tmp_path)
    assert md.endswith("\n")


if __name__ == "__main__":
    raise SystemExit(pytest.main([__file__, "-v"]))
