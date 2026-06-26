"""Generate synthetic .docx fixtures covering every case in plan.md section 5.

Run directly to (re)create the sample files in tests/fixtures/:

    python tests/generate_fixtures.py

The same builder functions are imported by test_converter.py so the test suite
does not depend on pre-generated files.
"""

from __future__ import annotations

import os
import struct
import tempfile
import zlib

from docx import Document
from docx.shared import Pt


def _make_png():
    """Build a valid 1x1 red PNG (with correct CRCs) without external assets."""
    def chunk(type_bytes, data):
        body = type_bytes + data
        crc = zlib.crc32(body) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + body + struct.pack(">I", crc)

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)  # 1x1, 8-bit, RGB
    raw = b"\x00\xff\x00\x00"  # one scanline: filter=0 + red pixel
    idat = zlib.compress(raw)
    return (signature + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


# A minimal valid PNG so fixtures can embed a real image without binary assets.
_PNG_1x1 = _make_png()

FIXTURE_DIR = os.path.join(os.path.dirname(__file__), "fixtures")


def _add_hyperlink(paragraph, url, text):
    """Add a real external hyperlink run to *paragraph* (python-docx has no API)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def build_inline_formatting():
    """Case 1: bold, italic, bold+italic, underline, strikethrough."""
    doc = Document()
    doc.add_heading("Inline Formatting", level=1)
    p = doc.add_paragraph()
    p.add_run("This is ")
    p.add_run("bold").bold = True
    p.add_run(", ")
    p.add_run("italic").italic = True
    p.add_run(", ")
    r = p.add_run("bold italic")
    r.bold = True
    r.italic = True
    p.add_run(", ")
    p.add_run("underlined").underline = True
    p.add_run(", and ")
    p.add_run("struck out").font.strike = True
    p.add_run(".")
    return doc


def build_headings():
    """Case 2: H1-H4 headings interleaved with normal paragraphs."""
    doc = Document()
    for level in range(1, 5):
        doc.add_heading("Heading Level %d" % level, level=level)
        doc.add_paragraph("Body text under heading %d." % level)
    return doc


def build_lists():
    """Case 3: bulleted list with nesting plus a numbered list."""
    doc = Document()
    doc.add_heading("Lists", level=1)
    doc.add_paragraph("First bullet", style="List Bullet")
    doc.add_paragraph("Nested bullet", style="List Bullet 2")
    doc.add_paragraph("Second bullet", style="List Bullet")
    doc.add_paragraph("Step one", style="List Number")
    doc.add_paragraph("Step two", style="List Number")
    return doc


def build_simple_table():
    """Case 4: a plain table with no merged cells."""
    doc = Document()
    doc.add_heading("Simple Table", level=1)
    table = doc.add_table(rows=2, cols=3)
    headers = ["Name", "Role", "City"]
    data = ["Ada", "Engineer", "London"]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    for i, text in enumerate(data):
        table.rows[1].cells[i].text = text
    return doc


def build_formatted_table():
    """Case 5: a table whose cells contain bold and italic text."""
    doc = Document()
    doc.add_heading("Formatted Table", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Feature").bold = True
    table.rows[0].cells[1].paragraphs[0].add_run("Status").bold = True
    table.rows[1].cells[0].paragraphs[0].add_run("Export").italic = True
    cell = table.rows[1].cells[1]
    cell.paragraphs[0].add_run("done").bold = True
    return doc


def build_merged_table():
    """Case 6: a table with a horizontally merged header (HTML fallback)."""
    doc = Document()
    doc.add_heading("Merged Table", level=1)
    table = doc.add_table(rows=3, cols=2)
    top = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    top.text = "Quarterly Results"
    table.rows[1].cells[0].text = "Q1"
    table.rows[1].cells[1].text = "100"
    table.rows[2].cells[0].text = "Q2"
    table.rows[2].cells[1].text = "150"
    return doc


def build_links_and_code():
    """Case 7: hyperlinks plus inline monospace (code) text."""
    doc = Document()
    doc.add_heading("Links and Code", level=1)
    p = doc.add_paragraph("Visit ")
    _add_hyperlink(p, "https://example.com", "Example")
    p.add_run(" for details.")
    p2 = doc.add_paragraph("Run ")
    code = p2.add_run("pip install docx")
    code.font.name = "Consolas"
    p2.add_run(" first.")
    return doc


def build_with_image():
    """Case 8: a document with one embedded image."""
    doc = Document()
    doc.add_heading("Image", level=1)
    doc.add_paragraph("Below is an embedded image:")
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    try:
        tmp.write(_PNG_1x1)
        tmp.close()
        doc.add_picture(tmp.name, width=Pt(48))
    finally:
        os.unlink(tmp.name)
    return doc


def build_kitchen_sink():
    """Case 9: a messy real-world-style document combining everything."""
    doc = Document()
    doc.add_heading("Project Report", level=1)
    doc.add_paragraph("An overview with a footnote-free summary.")
    doc.add_heading("Highlights", level=2)
    doc.add_paragraph("Shipped feature", style="List Bullet")
    doc.add_paragraph("Sub-detail", style="List Bullet 2")
    doc.add_paragraph("Next milestone", style="List Number")

    p = doc.add_paragraph()
    p.add_run("Important: ").bold = True
    p.add_run("see the ")
    _add_hyperlink(p, "https://example.com/spec", "spec")
    p.add_run(" and the ")
    p.add_run("config").font.name = "Courier New"
    p.add_run(" file.")

    doc.add_heading("Data", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Metric").bold = True
    table.rows[0].cells[1].paragraphs[0].add_run("Value").bold = True
    table.rows[1].cells[0].text = "Uptime"
    table.rows[1].cells[1].paragraphs[0].add_run("99.9%").italic = True

    doc.add_paragraph("End of report.")
    return doc


BUILDERS = {
    "01_inline_formatting": build_inline_formatting,
    "02_headings": build_headings,
    "03_lists": build_lists,
    "04_simple_table": build_simple_table,
    "05_formatted_table": build_formatted_table,
    "06_merged_table": build_merged_table,
    "07_links_and_code": build_links_and_code,
    "08_with_image": build_with_image,
    "09_kitchen_sink": build_kitchen_sink,
}


def generate_all(target_dir=FIXTURE_DIR):
    os.makedirs(target_dir, exist_ok=True)
    paths = {}
    for name, builder in BUILDERS.items():
        path = os.path.join(target_dir, name + ".docx")
        builder().save(path)
        paths[name] = path
    return paths


if __name__ == "__main__":
    written = generate_all()
    for name, path in written.items():
        print("wrote", path)
